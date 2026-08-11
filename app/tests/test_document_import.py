import os
import tempfile
import unittest
from unittest import mock

from docx import Document

import main
from main import read_imported_text


class DocumentImportTests(unittest.TestCase):
    def test_reads_utf8_and_cp1251_text(self):
        with tempfile.TemporaryDirectory() as folder:
            utf8_path = os.path.join(folder, "plan.txt")
            cp1251_path = os.path.join(folder, "old-plan.txt")
            with open(utf8_path, "w", encoding="utf-8-sig") as stream:
                stream.write("Тестовый план\nВторая строка")
            with open(cp1251_path, "wb") as stream:
                stream.write("Старый текст".encode("cp1251"))

            self.assertEqual(read_imported_text(utf8_path), "Тестовый план\nВторая строка")
            self.assertEqual(read_imported_text(cp1251_path), "Старый текст")

    def test_reads_docx_paragraphs_lists_and_tables(self):
        with tempfile.TemporaryDirectory() as folder:
            path = os.path.join(folder, "Тестовый план.docx")
            document = Document()
            document.add_heading("Тестовый план", level=1)
            document.add_paragraph("Первый поворот")
            document.add_paragraph("Второй поворот", style="List Bullet")
            table = document.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "Глава"
            table.cell(0, 1).text = "Событие"
            document.save(path)

            text = read_imported_text(path)

            self.assertIn("Тестовый план", text)
            self.assertIn("Первый поворот", text)
            self.assertIn("• Второй поворот", text)
            self.assertIn("Глава | Событие", text)

    def test_rejects_unsupported_document(self):
        with tempfile.TemporaryDirectory() as folder:
            path = os.path.join(folder, "plan.pdf")
            with open(path, "wb") as stream:
                stream.write(b"not a supported document")
            with self.assertRaises(ValueError):
                read_imported_text(path)

    def test_api_returns_text_and_filename(self):
        class FakeWebview:
            OPEN_DIALOG = 1

        class FakeWindow:
            def __init__(self, path):
                self.path = path

            def create_file_dialog(self, *args, **kwargs):
                return [self.path]

        with tempfile.TemporaryDirectory() as folder:
            path = os.path.join(folder, "Тестовый план.txt")
            with open(path, "w", encoding="utf-8") as stream:
                stream.write("Текст плана")
            with mock.patch.object(main, "webview", FakeWebview), mock.patch.object(main, "window", FakeWindow(path)):
                response = main.Api().import_text_document("ru")

            self.assertTrue(response["ok"])
            self.assertEqual(response["name"], "Тестовый план.txt")
            self.assertEqual(response["text"], "Текст плана")


if __name__ == "__main__":
    unittest.main()
