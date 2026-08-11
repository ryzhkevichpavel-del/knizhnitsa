# -*- coding: utf-8 -*-
"""
Книжница — нативное приложение для Windows для написания книг.
Окно на движке WebView2 (Edge), интерфейс — ui.html.
Данные хранятся в реальном файле: %APPDATA%\\Книжница\\library.json
"""
import os
import io
import sys
import re
import json
import time
import shutil
import base64
import ctypes
import mimetypes
import threading
import urllib.error
import urllib.parse
import urllib.request
import webbrowser
from datetime import datetime

from windows_startup import (
    APP_USER_MODEL_ID,
    MUTEX_NAME,
    SingleInstance,
    StartupLog,
    StartupSplash,
    activate_process_window,
    set_process_app_user_model_id,
)

APP_NAME = "Книжница"
APP_VERSION = "1.2.0"
LEGACY_APP_NAMES = ("Пиши книгу",)
MAX_BACKUPS = 15
AUTO_BACKUP_INTERVAL = 10 * 60
ALWAYS_START_MAXIMIZED = True
webview = None
window = None
startup_log = None
startup_splash = None
single_instance = None
last_auto_backup = 0
library_lock = threading.RLock()
UPDATE_API_URL = "https://api.github.com/repos/ryzhkevichpavel-del/knizhnitsa/releases/latest"


def data_dir():
    base = os.environ.get("APPDATA") or os.path.expanduser("~")
    d = os.path.join(base, APP_NAME)
    os.makedirs(d, exist_ok=True)
    return d


DATA_DIR = data_dir()
DATA_FILE = os.path.join(DATA_DIR, "library.json")
BACKUP_DIR = os.path.join(DATA_DIR, "Резервные копии")
WINDOW_STATE_FILE = os.path.join(DATA_DIR, "window.json")


def resource(name):
    base = getattr(sys, "_MEIPASS", os.path.dirname(os.path.abspath(__file__)))
    return os.path.join(base, name)


def appdata_dir(app_name):
    base = os.environ.get("APPDATA") or os.path.expanduser("~")
    return os.path.join(base, app_name)


def result(ok=True, error="", **extra):
    response = {"ok": bool(ok), "error": error}
    response.update(extra)
    return response


def norm_lang(lang):
    return "en" if lang == "en" else "ru"


def msg(lang, ru, en):
    return en if norm_lang(lang) == "en" else ru


def read_text(path):
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


def parse_library_text(raw):
    """Проверить, что строка является библиотекой, не меняя её содержимое."""
    data = json.loads(raw)
    if not isinstance(data, dict):
        raise ValueError("library root must be an object")
    return data


def read_valid_library(path):
    raw = read_text(path)
    return raw, parse_library_text(raw)


def atomic_write_text(path, data):
    """Записать текст через временный файл и принудительно сбросить его на диск."""
    os.makedirs(os.path.dirname(path), exist_ok=True)
    tmp = path + ".tmp"
    with open(tmp, "w", encoding="utf-8", newline="") as stream:
        stream.write(data)
        stream.flush()
        os.fsync(stream.fileno())
    os.replace(tmp, path)


def library_stats(data):
    """Небольшая диагностика размера библиотеки без раскрытия её содержимого."""
    if isinstance(data, str):
        raw = data
        parsed = parse_library_text(data)
    else:
        parsed = data if isinstance(data, dict) else {}
        raw = json.dumps(parsed, ensure_ascii=False)
    books = parsed.get("books") if isinstance(parsed.get("books"), list) else []
    chapters = characters = history = 0
    for book in books:
        if not isinstance(book, dict):
            continue
        chapter_list = book.get("chapters") if isinstance(book.get("chapters"), list) else []
        character_list = book.get("characters") if isinstance(book.get("characters"), list) else []
        chapters += len(chapter_list)
        characters += len(character_list)
        for collection_owner in [book, *chapter_list, *character_list]:
            if not isinstance(collection_owner, dict):
                continue
            for field in ("history", "versions", "planVersions", "arcVersions"):
                if isinstance(collection_owner.get(field), list):
                    history += len(collection_owner[field])
    return {
        "bytes": len(raw.encode("utf-8")),
        "books": len(books),
        "chapters": chapters,
        "characters": characters,
        "historyEntries": history,
    }


def _open_clipboard(attempts=8, delay=0.02):
    """Открыть системный буфер обмена, переждав короткую блокировку другим приложением."""
    user32 = ctypes.windll.user32
    user32.OpenClipboard.argtypes = [ctypes.c_void_p]
    user32.OpenClipboard.restype = ctypes.c_int
    for _ in range(attempts):
        if user32.OpenClipboard(None):
            return True
        time.sleep(delay)
    return False


def read_clipboard_text():
    """Прочитать обычный Unicode-текст из буфера обмена Windows."""
    cf_unicode_text = 13
    user32 = ctypes.windll.user32
    kernel32 = ctypes.windll.kernel32
    user32.IsClipboardFormatAvailable.argtypes = [ctypes.c_uint]
    user32.IsClipboardFormatAvailable.restype = ctypes.c_int
    user32.GetClipboardData.argtypes = [ctypes.c_uint]
    user32.GetClipboardData.restype = ctypes.c_void_p
    kernel32.GlobalLock.argtypes = [ctypes.c_void_p]
    kernel32.GlobalLock.restype = ctypes.c_void_p
    kernel32.GlobalUnlock.argtypes = [ctypes.c_void_p]
    kernel32.GlobalUnlock.restype = ctypes.c_int

    if not _open_clipboard():
        raise OSError("clipboard is busy")
    handle = None
    pointer = None
    try:
        if not user32.IsClipboardFormatAvailable(cf_unicode_text):
            return ""
        handle = user32.GetClipboardData(cf_unicode_text)
        if not handle:
            raise OSError("clipboard text is unavailable")
        pointer = kernel32.GlobalLock(handle)
        if not pointer:
            raise OSError("clipboard text is locked")
        return ctypes.wstring_at(pointer)
    finally:
        if pointer and handle:
            kernel32.GlobalUnlock(handle)
        user32.CloseClipboard()


def write_clipboard_text(text):
    """Записать обычный Unicode-текст в буфер обмена Windows."""
    cf_unicode_text = 13
    gmem_moveable = 0x0002
    data = str(text or "").encode("utf-16-le") + b"\x00\x00"
    user32 = ctypes.windll.user32
    kernel32 = ctypes.windll.kernel32
    kernel32.GlobalAlloc.argtypes = [ctypes.c_uint, ctypes.c_size_t]
    kernel32.GlobalAlloc.restype = ctypes.c_void_p
    kernel32.GlobalLock.argtypes = [ctypes.c_void_p]
    kernel32.GlobalLock.restype = ctypes.c_void_p
    kernel32.GlobalUnlock.argtypes = [ctypes.c_void_p]
    kernel32.GlobalUnlock.restype = ctypes.c_int
    kernel32.GlobalFree.argtypes = [ctypes.c_void_p]
    kernel32.GlobalFree.restype = ctypes.c_void_p
    user32.SetClipboardData.argtypes = [ctypes.c_uint, ctypes.c_void_p]
    user32.SetClipboardData.restype = ctypes.c_void_p

    handle = kernel32.GlobalAlloc(gmem_moveable, len(data))
    if not handle:
        raise OSError("clipboard allocation failed")
    pointer = kernel32.GlobalLock(handle)
    if not pointer:
        kernel32.GlobalFree(handle)
        raise OSError("clipboard allocation is locked")
    ctypes.memmove(pointer, data, len(data))
    kernel32.GlobalUnlock(handle)

    if not _open_clipboard():
        kernel32.GlobalFree(handle)
        raise OSError("clipboard is busy")
    transferred = False
    try:
        if not user32.EmptyClipboard():
            raise OSError("clipboard could not be cleared")
        if not user32.SetClipboardData(cf_unicode_text, handle):
            raise OSError("clipboard text could not be written")
        transferred = True
    finally:
        user32.CloseClipboard()
        if not transferred:
            kernel32.GlobalFree(handle)


def read_imported_text(path):
    """Извлечь обычный текст из поддерживаемого авторского документа."""
    ext = os.path.splitext(path)[1].lower()
    if ext == ".docx":
        from docx import Document
        from docx.document import Document as DocumentObject
        from docx.oxml.table import CT_Tbl
        from docx.oxml.text.paragraph import CT_P
        from docx.table import Table
        from docx.text.paragraph import Paragraph

        document = Document(path)
        blocks = []
        parent = document.element.body if isinstance(document, DocumentObject) else document._element
        for child in parent.iterchildren():
            if isinstance(child, CT_P):
                paragraph = Paragraph(child, document)
                text = paragraph.text.strip()
                if not text:
                    continue
                style_name = (paragraph.style.name if paragraph.style is not None else "").lower()
                numbering = (
                    (paragraph._p.pPr is not None and paragraph._p.pPr.numPr is not None)
                    or "list" in style_name
                    or "спис" in style_name
                )
                blocks.append(("• " if numbering else "") + text)
            elif isinstance(child, CT_Tbl):
                table = Table(child, document)
                rows = []
                for row in table.rows:
                    cells = [re.sub(r"\s+", " ", cell.text).strip() for cell in row.cells]
                    if any(cells):
                        rows.append(" | ".join(cells))
                if rows:
                    blocks.append("\n".join(rows))
        return "\n\n".join(blocks).strip()

    if ext not in (".txt", ".md"):
        raise ValueError("unsupported document type")
    with open(path, "rb") as stream:
        raw = stream.read()
    for encoding in ("utf-8-sig", "utf-8", "cp1251"):
        try:
            return raw.decode(encoding).replace("\r\n", "\n").replace("\r", "\n")
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace").replace("\r\n", "\n").replace("\r", "\n")


def build_book_txt(book, lang="ru"):
    """Собрать понятный текстовый файл книги с заголовками глав."""
    lang = norm_lang(lang)
    book = book if isinstance(book, dict) else {}
    title = book.get("title") or msg(lang, "Книга", "Book")
    parts = [str(title).strip()]
    author = str(book.get("author") or "").strip()
    if author:
        parts.append(author)
    for index, chapter in enumerate(book.get("chapters") or []):
        if not isinstance(chapter, dict):
            continue
        heading = chapter.get("title") or f"{msg(lang, 'Глава', 'Chapter')} {index + 1}"
        content = str(chapter.get("content") or "").replace("\r\n", "\n").replace("\r", "\n").strip()
        parts.append(str(heading).strip() + ("\n\n" + content if content else ""))
    return "\n\n".join(parts).rstrip() + "\n"


def write_book_docx(book, path, lang="ru"):
    """Записать книгу в DOCX, сохранив абзацы и одиночные переносы строк."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, Inches

    lang = norm_lang(lang)
    book = book if isinstance(book, dict) else {}
    title = book.get("title") or msg(lang, "книга", "book")
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    normal = doc.styles["Normal"]
    normal.font.name = "Georgia"
    normal.font.size = Pt(12)

    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(str(title))
    title_run.bold = True
    title_run.font.size = Pt(22)

    author = str(book.get("author") or "").strip()
    if author:
        author_paragraph = doc.add_paragraph()
        author_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_run = author_paragraph.add_run(author)
        author_run.font.size = Pt(12)

    chapters = book.get("chapters") or []
    for index, chapter in enumerate(chapters):
        if not isinstance(chapter, dict):
            continue
        doc.add_page_break()
        heading = doc.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading_run = heading.add_run(chapter.get("title") or f"{msg(lang, 'Глава', 'Chapter')} {index + 1}")
        heading_run.bold = True
        heading_run.font.size = Pt(16)

        text = str(chapter.get("content") or "").replace("\r\n", "\n").replace("\r", "\n")
        for block in re.split(r"\n[ \t]*\n", text):
            if not block.strip():
                continue
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.first_line_indent = Inches(0.35)
            paragraph.paragraph_format.line_spacing = 1.25
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            lines = block.strip().split("\n")
            for line_index, line in enumerate(lines):
                if line_index:
                    paragraph.add_run().add_break()
                paragraph.add_run(line)
    doc.save(path)


def has_books(raw):
    try:
        data = parse_library_text(raw)
        return bool(data.get("books"))
    except Exception:
        return False


def current_file_has_books():
    if not os.path.exists(DATA_FILE):
        return False
    try:
        return has_books(read_text(DATA_FILE))
    except Exception:
        return False


def migrate_legacy_data():
    current_raw = ""
    if os.path.exists(DATA_FILE):
        try:
            current_raw, _ = read_valid_library(DATA_FILE)
            # Любой корректный текущий файл авторитетен, даже если пользователь
            # сознательно удалил из него все книги.
            return result(True, migrated=False)
        except Exception as exc:
            return result(
                False,
                f"Основной файл библиотеки повреждён; перенос старой версии остановлен: {exc}",
                code="current_file_invalid",
                migrated=False,
            )
    errors = []
    for old_name in LEGACY_APP_NAMES:
        old_file = os.path.join(appdata_dir(old_name), "library.json")
        if not os.path.exists(old_file):
            continue
        try:
            raw = read_text(old_file)
            if not has_books(raw):
                continue
            os.makedirs(os.path.dirname(DATA_FILE), exist_ok=True)
            if current_raw:
                backup = create_backup_from_text(current_raw, "перед переносом")
                if not backup.get("ok"):
                    return result(False, backup.get("error", ""), migrated=False)
            shutil.copy2(old_file, DATA_FILE)
            create_backup_from_text(raw, "перенесено из старой версии")
            return result(True, migrated=True, source=old_file)
        except Exception as exc:
            errors.append(f"{old_file}: {exc}")
            continue
    if errors:
        return result(False, "Не удалось перенести старую библиотеку: " + "; ".join(errors), migrated=False)
    return result(True, migrated=False)


def safe_slug(text):
    text = re.sub(r'[<>:"/\\|?*\x00-\x1f]+', " ", text or "").strip()
    text = re.sub(r"\s+", " ", text)
    return text[:70] or APP_NAME


def backup_name(reason):
    stamp = datetime.now().strftime("%Y-%m-%d %H-%M-%S-%f")
    return f"{stamp} - {safe_slug(reason)}.json"


def backup_files():
    try:
        files = [
            os.path.join(BACKUP_DIR, name)
            for name in os.listdir(BACKUP_DIR)
            if name.lower().endswith(".json")
        ]
        return sorted(files, key=os.path.getmtime, reverse=True)
    except (FileNotFoundError, NotADirectoryError):
        return []


def valid_recovery_candidates(include_tmp=True):
    """Вернуть только читаемые JSON-копии, начиная с самой свежей."""
    candidates = []
    tmp = DATA_FILE + ".tmp"
    paths = ([tmp] if include_tmp and os.path.exists(tmp) else []) + backup_files()
    for path in paths:
        try:
            raw, _ = read_valid_library(path)
            candidates.append({
                "path": os.path.abspath(path),
                "source": "temporary" if os.path.abspath(path) == os.path.abspath(tmp) else "backup",
                "modified": os.path.getmtime(path),
                "data": raw,
            })
        except Exception:
            continue
    return sorted(candidates, key=lambda item: item["modified"], reverse=True)


def latest_recovery_candidate(include_tmp=True):
    candidates = valid_recovery_candidates(include_tmp)
    return candidates[0] if candidates else None


def read_latest_backup_bytes():
    for candidate in valid_recovery_candidates(include_tmp=False):
        return candidate["data"].encode("utf-8")
    return None


def recycle_path(path):
    try:
        from ctypes import wintypes

        class SHFILEOPSTRUCTW(ctypes.Structure):
            _fields_ = [
                ("hwnd", wintypes.HWND),
                ("wFunc", wintypes.UINT),
                ("pFrom", wintypes.LPCWSTR),
                ("pTo", wintypes.LPCWSTR),
                ("fFlags", wintypes.WORD),
                ("fAnyOperationsAborted", wintypes.BOOL),
                ("hNameMappings", wintypes.LPVOID),
                ("lpszProgressTitle", wintypes.LPCWSTR),
            ]

        op = SHFILEOPSTRUCTW()
        op.wFunc = 3
        op.pFrom = path + "\0\0"
        op.fFlags = 0x0040 | 0x0010 | 0x0400 | 0x0004
        if ctypes.windll.shell32.SHFileOperationW(ctypes.byref(op)) == 0:
            return True
    except Exception:
        pass
    try:
        os.remove(path)
        return not os.path.exists(path)
    except Exception:
        pass
    return False


def prune_backups():
    removed = 0
    failed = []
    try:
        for path in backup_files()[MAX_BACKUPS:]:
            if recycle_path(path):
                removed += 1
            else:
                failed.append(path)
        return result(not failed, "" if not failed else "Не удалось убрать старые резервные копии.", removed=removed)
    except Exception as exc:
        return result(False, f"Не удалось проверить резервные копии: {exc}", removed=removed)


def create_backup_from_text(data, reason, lang="ru"):
    if not data:
        return result(True)
    try:
        parse_library_text(data)
        raw = data.encode("utf-8")
        if read_latest_backup_bytes() == raw:
            return result(True)
        os.makedirs(BACKUP_DIR, exist_ok=True)
        path = os.path.join(BACKUP_DIR, backup_name(reason))
        atomic_write_text(path, data)
        prune_result = prune_backups()
        if not prune_result.get("ok"):
            return result(False, prune_result.get("error", ""), path=path)
        return {"ok": True, "path": path, "error": ""}
    except Exception as e:
        return result(False, f"{msg(lang, 'Не удалось сделать резервную копию', 'Could not create a backup')}: {e}")


def backup_current_file(reason, force=False, lang="ru"):
    global last_auto_backup
    if not os.path.exists(DATA_FILE):
        return result(True)
    if not force and time.time() - last_auto_backup < AUTO_BACKUP_INTERVAL:
        return result(True)
    try:
        res = create_backup_from_text(read_text(DATA_FILE), reason, lang)
        if res.get("ok"):
            last_auto_backup = time.time()
        return res
    except Exception as e:
        return result(False, f"{msg(lang, 'Не удалось сделать резервную копию', 'Could not create a backup')}: {e}")


def load_initial_settings():
    settings = {"theme": "light", "accent": "#3f4ea3", "fontSize": 19, "lang": "ru"}
    try:
        if os.path.exists(DATA_FILE):
            data = json.loads(read_text(DATA_FILE))
            saved = data.get("settings") or {}
            if saved.get("theme") in ("light", "dark"):
                settings["theme"] = saved["theme"]
            if re.match(r"^#[0-9a-fA-F]{6}$", saved.get("accent") or ""):
                settings["accent"] = saved["accent"]
            if isinstance(saved.get("fontSize"), int):
                settings["fontSize"] = saved["fontSize"]
            if saved.get("lang") == "en":
                settings["lang"] = "en"
    except Exception:
        pass
    return settings


def accent_soft(hex_color, theme):
    h = hex_color.lstrip("#")
    r, g, b = int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)
    alpha = 0.22 if theme == "dark" else 0.13
    return f"rgba({r},{g},{b},{alpha})"


def prepare_html(html, settings):
    theme = settings["theme"]
    accent = settings["accent"]
    lang = norm_lang(settings.get("lang"))
    soft = accent_soft(accent, theme)
    html = html.replace('data-theme="light"', f'data-theme="{theme}"', 1)
    html = re.sub(r'<html lang="[^"]+"', f'<html lang="{lang}"', html, count=1)
    initial = (
        "<script>"
        f"window.__INITIAL_SETTINGS__={json.dumps(settings, ensure_ascii=False)};"
        "</script>"
    )
    style = (
        "<style id=\"initialAccent\">"
        f":root,[data-theme=\"light\"],[data-theme=\"dark\"]{{--accent:{accent};--accent-soft:{soft};}}"
        "</style>"
    )
    return html.replace("</head>", initial + style + "</head>", 1)


def load_window_state():
    try:
        with open(WINDOW_STATE_FILE, "r", encoding="utf-8") as f:
            s = json.load(f)
        w = int(s.get("width", 0))
        h = int(s.get("height", 0))
        if w < 940 or h < 620:          # не меньше min_size
            return None
        state = {"width": w, "height": h}
        if isinstance(s.get("x"), int) and isinstance(s.get("y"), int):
            state["x"] = s["x"]
            state["y"] = s["y"]
        state["maximized"] = bool(s.get("maximized"))
        return state
    except Exception:
        return None


def save_window_state():
    if window is None:
        return
    try:
        state = {
            "width": int(window.width),
            "height": int(window.height),
            "x": int(window.x),
            "y": int(window.y),
        }
        tmp = WINDOW_STATE_FILE + ".tmp"
        with open(tmp, "w", encoding="utf-8") as f:
            json.dump(state, f)
        os.replace(tmp, WINDOW_STATE_FILE)
    except Exception:
        pass


class Api:
    """Мост между интерфейсом (JS) и диском (Python)."""

    # ---- данные книг ----
    def load(self):
        """Загрузить библиотеку и никогда не маскировать ошибку под пустую книгу."""
        with library_lock:
            try:
                raw, parsed = read_valid_library(DATA_FILE)
                recovery = latest_recovery_candidate()
                main_modified = os.path.getmtime(DATA_FILE)
                newer_tmp = bool(
                    recovery
                    and recovery["source"] == "temporary"
                    and recovery["modified"] > main_modified
                    and recovery["data"] != raw
                )
                return result(
                    True,
                    data=raw,
                    stats=library_stats(raw),
                    recovery_available=newer_tmp,
                    recovery_source=recovery["source"] if newer_tmp else "",
                    recovery_path=recovery["path"] if newer_tmp else "",
                )
            except FileNotFoundError:
                recovery = latest_recovery_candidate()
                if recovery:
                    return result(
                        False,
                        "Основной файл библиотеки не найден, но доступна резервная копия.",
                        code="library_missing",
                        data="",
                        recovery_available=True,
                        recovery_source=recovery["source"],
                        recovery_path=recovery["path"],
                    )
                return result(True, data="", stats=library_stats({}), recovery_available=False)
            except Exception as exc:
                recovery = latest_recovery_candidate()
                return result(
                    False,
                    f"Не удалось прочитать библиотеку: {exc}",
                    code="library_invalid",
                    data="",
                    recovery_available=bool(recovery),
                    recovery_source=(recovery or {}).get("source", ""),
                    recovery_path=(recovery or {}).get("path", ""),
                )

    def save(self, data, lang="ru"):
        return self._save_library(data, lang)

    def flush_save(self, data, lang="ru"):
        """Синхронная точка сохранения для обработчика закрытия окна."""
        return self._save_library(data, lang)

    def _save_library(self, data, lang="ru"):
        lang = norm_lang(lang)
        with library_lock:
            try:
                if not isinstance(data, str):
                    raise ValueError(msg(lang, "Данные библиотеки должны быть текстом JSON.", "Library data must be JSON text."))
                parsed = parse_library_text(data)
                old = ""
                if os.path.exists(DATA_FILE):
                    try:
                        old, _ = read_valid_library(DATA_FILE)
                    except Exception as exc:
                        recovery = latest_recovery_candidate()
                        return result(
                            False,
                            f"{msg(lang, 'Основной файл повреждён; сначала восстановите копию', 'The main file is damaged; restore a backup first')}: {exc}",
                            code="current_file_invalid",
                            recovery_available=bool(recovery),
                            recovery_path=(recovery or {}).get("path", ""),
                        )
                if old == data:
                    return result(True, changed=False, stats=library_stats(data))
                if old:
                    backup_res = backup_current_file(
                        msg(lang, "автокопия перед сохранением", "auto backup before saving"),
                        lang=lang,
                    )
                    if not backup_res.get("ok"):
                        return backup_res
                atomic_write_text(DATA_FILE, data)
                return result(True, changed=True, stats=library_stats(data))
            except Exception as e:
                return result(False, f"{msg(lang, 'Не удалось сохранить библиотеку', 'Could not save library')}: {e}")

    def recover_latest_backup(self, lang="ru"):
        """Безопасно восстановить свежую исправную копию и вернуть её интерфейсу."""
        lang = norm_lang(lang)
        with library_lock:
            recovery = latest_recovery_candidate()
            if not recovery:
                return result(False, msg(lang, "Исправная резервная копия не найдена.", "No valid backup was found."))
            try:
                if os.path.exists(DATA_FILE):
                    old = read_text(DATA_FILE)
                    if old and old != recovery["data"]:
                        try:
                            parse_library_text(old)
                            backup = create_backup_from_text(
                                old,
                                msg(lang, "перед восстановлением", "before recovery"),
                                lang,
                            )
                            if not backup.get("ok"):
                                return backup
                        except Exception:
                            os.makedirs(BACKUP_DIR, exist_ok=True)
                            damaged_path = os.path.join(BACKUP_DIR, backup_name(msg(lang, "поврежденный файл", "damaged file")))
                            atomic_write_text(damaged_path, old)
                            prune_backups()
                atomic_write_text(DATA_FILE, recovery["data"])
                return result(
                    True,
                    data=recovery["data"],
                    source=recovery["source"],
                    path=recovery["path"],
                    restored=True,
                    stats=library_stats(recovery["data"]),
                )
            except Exception as exc:
                return result(False, f"{msg(lang, 'Не удалось восстановить библиотеку', 'Could not restore the library')}: {exc}")

    def data_stats(self, data=None, lang="ru"):
        try:
            if data is None:
                data = read_text(DATA_FILE) if os.path.exists(DATA_FILE) else "{}"
            return result(True, stats=library_stats(data))
        except Exception as exc:
            return result(False, f"{msg(lang, 'Не удалось проверить библиотеку', 'Could not inspect the library')}: {exc}")

    def maintain_data(self, lang="ru"):
        """Безопасная уборка старых копий и заведомо ненужного старого tmp."""
        with library_lock:
            prune = prune_backups()
            tmp_removed = False
            tmp = DATA_FILE + ".tmp"
            try:
                if os.path.exists(tmp) and os.path.exists(DATA_FILE) and os.path.getmtime(tmp) <= os.path.getmtime(DATA_FILE):
                    os.remove(tmp)
                    tmp_removed = True
            except OSError:
                pass
            return result(prune.get("ok", False), prune.get("error", ""), removed=prune.get("removed", 0), tempRemoved=tmp_removed)

    def data_location(self):
        return DATA_FILE

    def app_info(self):
        return result(
            True,
            version=APP_VERSION,
            data_location=DATA_FILE,
            startup_log=os.path.join(DATA_DIR, "startup.log"),
        )

    def check_for_updates(self, lang="ru"):
        """Проверить выпуск только по явному вызову из настроек."""
        lang = norm_lang(lang)
        try:
            request = urllib.request.Request(
                UPDATE_API_URL,
                headers={"Accept": "application/vnd.github+json", "User-Agent": f"Knizhnitsa/{APP_VERSION}"},
            )

            with urllib.request.urlopen(request, timeout=8) as response:
                payload = json.loads(response.read().decode("utf-8"))
            latest = str(payload.get("tag_name") or payload.get("name") or "").strip().lstrip("vV")
            if not latest:
                raise ValueError("release version is missing")
            url = str(payload.get("html_url") or "https://github.com/ryzhkevichpavel-del/knizhnitsa/releases/latest")

            def version_tuple(value):
                numbers = [int(part) for part in re.findall(r"\d+", value)[:4]]
                return tuple(numbers + [0] * (4 - len(numbers)))

            return result(
                True,
                current_version=APP_VERSION,
                latest_version=latest,
                update_available=version_tuple(latest) > version_tuple(APP_VERSION),
                url=url,
            )
        except Exception as exc:
            return result(
                False,
                f"{msg(lang, 'Не удалось проверить обновления', 'Could not check for updates')}: {exc}",
                current_version=APP_VERSION,
                latest_version="",
                update_available=False,
                url="",
            )

    def open_external(self, url, lang="ru"):
        """Открыть только страницу GitHub по явному клику пользователя."""
        lang = norm_lang(lang)
        try:
            parsed = urllib.parse.urlparse(str(url or ""))
            if parsed.scheme != "https" or parsed.hostname not in {"github.com", "www.github.com"}:
                raise ValueError(msg(lang, "Разрешены только ссылки GitHub по HTTPS.", "Only HTTPS GitHub links are allowed."))
            opened = webbrowser.open(parsed.geturl(), new=2)
            return result(bool(opened), "" if opened else msg(lang, "Не удалось открыть браузер.", "Could not open the browser."))
        except Exception as exc:
            return result(False, f"{msg(lang, 'Не удалось открыть ссылку', 'Could not open the link')}: {exc}")

    # ---- системный буфер обмена для меню правой кнопки ----
    def clipboard_read_text(self, lang="ru"):
        lang = norm_lang(lang)
        try:
            return {"ok": True, "error": "", "text": read_clipboard_text()}
        except Exception:
            return result(False, msg(lang, "Не удалось прочитать буфер обмена.", "Could not read the clipboard."))

    def clipboard_write_text(self, text, lang="ru"):
        lang = norm_lang(lang)
        try:
            write_clipboard_text(text)
            return result(True)
        except Exception:
            return result(False, msg(lang, "Не удалось записать текст в буфер обмена.", "Could not write to the clipboard."))

    def backup_state(self, data, reason="ручная копия", lang="ru"):
        return create_backup_from_text(data, reason, lang)

    # ---- цвет системного заголовка окна (Windows 11) ----
    def set_titlebar(self, caption_hex):
        try:
            import ctypes
            user32 = ctypes.windll.user32
            user32.FindWindowW.restype = ctypes.c_void_p
            user32.FindWindowW.argtypes = [ctypes.c_wchar_p, ctypes.c_wchar_p]
            hwnd = user32.FindWindowW(None, APP_NAME)
            if not hwnd:
                return False
            dwm = ctypes.windll.dwmapi.DwmSetWindowAttribute
            dwm.argtypes = [ctypes.c_void_p, ctypes.c_uint, ctypes.c_void_p, ctypes.c_uint]

            def cref(h):
                h = h.lstrip("#")
                r, g, b = int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)
                return r | (g << 8) | (b << 16)  # COLORREF = 0x00BBGGRR

            dark = ctypes.c_int(1)
            dwm(hwnd, 20, ctypes.byref(dark), 4)                       # IMMERSIVE_DARK_MODE → светлые кнопки
            cap = ctypes.c_int(cref(caption_hex))
            dwm(hwnd, 35, ctypes.byref(cap), 4)                        # CAPTION_COLOR
            txt = ctypes.c_int(cref("#ffffff"))
            dwm(hwnd, 36, ctypes.byref(txt), 4)                        # TEXT_COLOR
            return True
        except Exception:
            return False

    # ---- выбор изображения (фото персонажа / фон карты) ----
    def pick_image(self, lang="ru"):
        lang = norm_lang(lang)
        try:
            paths = window.create_file_dialog(
                webview.OPEN_DIALOG,
                file_types=(
                    msg(lang, "Изображения (*.png;*.jpg;*.jpeg;*.webp;*.bmp;*.gif)", "Images (*.png;*.jpg;*.jpeg;*.webp;*.bmp;*.gif)"),
                    msg(lang, "Все файлы (*.*)", "All files (*.*)"),
                ),
            )
            if not paths:
                return ""
            p = paths[0] if isinstance(paths, (list, tuple)) else paths
            # уменьшаем картинку, чтобы файл библиотеки не разрастался
            try:
                from PIL import Image
                img = Image.open(p)
                img.thumbnail((1000, 1000))
                buf = io.BytesIO()
                if img.mode in ("RGBA", "LA", "P"):
                    img = img.convert("RGBA")
                    img.save(buf, "PNG")
                    mime = "image/png"
                else:
                    img = img.convert("RGB")
                    img.save(buf, "JPEG", quality=84)
                    mime = "image/jpeg"
                data = base64.b64encode(buf.getvalue()).decode("ascii")
                return "data:%s;base64,%s" % (mime, data)
            except Exception:
                with open(p, "rb") as f:
                    data = base64.b64encode(f.read()).decode("ascii")
                mime = mimetypes.guess_type(p)[0] or "image/png"
                return "data:%s;base64,%s" % (mime, data)
        except Exception:
            return ""

    # ---- импорт текста в открытую главу / план / арку ----
    def import_text_document(self, lang="ru"):
        lang = norm_lang(lang)
        try:
            paths = window.create_file_dialog(
                webview.OPEN_DIALOG,
                file_types=(
                    msg(lang, "Документы Word и текст (*.docx;*.txt;*.md)", "Word and text documents (*.docx;*.txt;*.md)"),
                    msg(lang, "Документ Word (*.docx)", "Word document (*.docx)"),
                    msg(lang, "Текстовый документ (*.txt;*.md)", "Text document (*.txt;*.md)"),
                ),
            )
            if not paths:
                return {"ok": False, "cancelled": True, "error": ""}
            path = paths[0] if isinstance(paths, (list, tuple)) else paths
            text = read_imported_text(path)
            return {
                "ok": True,
                "error": "",
                "name": os.path.basename(path),
                "text": text,
            }
        except ValueError:
            return result(False, msg(lang, "Поддерживаются только DOCX, TXT и MD.", "Only DOCX, TXT, and MD are supported."))
        except Exception as exc:
            return result(False, f"{msg(lang, 'Не удалось прочитать документ', 'Could not read the document')}: {exc}")

    # ---- экспорт готовой книги ----
    def export_docx(self, book, lang="ru"):
        lang = norm_lang(lang)
        try:
            title = (book or {}).get("title") or msg(lang, "книга", "book")
            path = window.create_file_dialog(
                webview.SAVE_DIALOG,
                save_filename=f"{safe_slug(title)}.docx",
                file_types=(msg(lang, "Документ Word (*.docx)", "Word document (*.docx)"),),
            )
            if not path:
                return result(False, "")
            if isinstance(path, (list, tuple)):
                path = path[0]

            write_book_docx(book, path, lang)
            return result(True)
        except Exception as e:
            return result(False, f"{msg(lang, 'Не удалось сохранить DOCX', 'Could not save DOCX')}: {e}")

    def export_txt(self, content, lang="ru"):
        lang = norm_lang(lang)
        if isinstance(content, dict):
            content = build_book_txt(content, lang)
        return self._write_dialog(content, msg(lang, "книга.txt", "book.txt"),
                                  (msg(lang, "Текстовый файл (*.txt)", "Text file (*.txt)"),), encoding="utf-8", lang=lang)

    def _write_dialog(self, content, fname, ftypes, encoding="utf-8", lang="ru"):
        lang = norm_lang(lang)
        try:
            path = window.create_file_dialog(
                webview.SAVE_DIALOG, save_filename=fname, file_types=ftypes)
            if not path:
                return False
            if isinstance(path, (list, tuple)):
                path = path[0]
            with open(path, "w", encoding=encoding) as f:
                f.write(content)
            return result(True)
        except Exception as e:
            return result(False, f"{msg(lang, 'Не удалось сохранить файл', 'Could not save file')}: {e}")

    # ---- резервные копии ----
    def export_backup(self, data, lang="ru"):
        lang = norm_lang(lang)
        return self._write_dialog(
            data,
            msg(lang, "Книжница-копия.json", "Knizhnitsa-backup.json"),
            (msg(lang, "Файл копии (*.json)", "Backup file (*.json)"),),
            lang=lang,
        )

    def import_backup(self, lang="ru"):
        lang = norm_lang(lang)
        try:
            paths = window.create_file_dialog(
                webview.OPEN_DIALOG, file_types=(msg(lang, "Файл копии (*.json)", "Backup file (*.json)"),))
            if not paths:
                return result(False, "", cancelled=True)
            p = paths[0] if isinstance(paths, (list, tuple)) else paths
            raw, parsed = read_valid_library(p)
            return result(True, data=raw, name=os.path.basename(p), stats=library_stats(raw), cancelled=False)
        except (json.JSONDecodeError, ValueError) as exc:
            return result(False, f"{msg(lang, 'Файл копии повреждён или имеет неверный формат', 'The backup is damaged or has an invalid format')}: {exc}", cancelled=False)
        except Exception as exc:
            return result(False, f"{msg(lang, 'Не удалось прочитать файл копии', 'Could not read the backup file')}: {exc}", cancelled=False)


def main():
    global webview, window, startup_log, startup_splash, single_instance

    startup_log = StartupLog(DATA_DIR)
    app_id_ok = set_process_app_user_model_id(APP_USER_MODEL_ID)
    single_instance = SingleInstance(DATA_DIR, startup_log, mutex_name=MUTEX_NAME)
    role = single_instance.acquire()

    if role == "secondary":
        return
    if role == "error":
        ctypes.windll.user32.MessageBoxW(
            None,
            "Не удалось проверить уже запущенную Книжницу. Попробуйте ещё раз.",
            APP_NAME,
            0x10,
        )
        return

    startup_log.rotate_if_needed()
    startup_log.write(
        "python_entry",
        version=APP_VERSION,
        frozen=bool(getattr(sys, "frozen", False)),
        app_user_model_id=app_id_ok,
    )
    startup_log.write("instance_primary")

    window_ready = threading.Event()
    activation_pending = threading.Event()

    def request_window_activation():
        activation_pending.set()
        if window_ready.is_set():
            activate_process_window(os.getpid(), APP_NAME, logger=startup_log)

    single_instance.start_listener(request_window_activation)

    startup_splash = StartupSplash(resource("icon.ico"))
    splash_ok = startup_splash.start()
    startup_log.write("splash_shown", success=splash_ok)

    try:
        startup_log.write("webview_import_begin")
        import webview as webview_module

        webview = webview_module
        startup_log.write("webview_import_done")

        startup_log.write("data_migration_begin")
        migration_result = migrate_legacy_data()
        startup_log.write(
            "data_migration_done",
            success=bool(migration_result.get("ok")),
            migrated=bool(migration_result.get("migrated")),
            error=migration_result.get("error", ""),
        )

        settings = load_initial_settings()
        startup_log.write("html_prepare_begin")
        with open(resource("ui.html"), "r", encoding="utf-8") as f:
            html = prepare_html(f.read(), settings)
        startup_log.write("html_prepare_done")

        state = load_window_state()
        api = Api()
        win_kwargs = dict(
            html=html,
            js_api=api,
            width=(state or {}).get("width", 1240),
            height=(state or {}).get("height", 860),
            min_size=(940, 620),
            background_color="#161618" if settings["theme"] == "dark" else "#f6f6f4",
        )
        if state and "x" in state and "y" in state:
            win_kwargs["x"] = state["x"]
            win_kwargs["y"] = state["y"]

        startup_log.write("create_window_begin")
        window = webview.create_window(APP_NAME, **win_kwargs)
        startup_log.write("create_window_done")

        def on_window_shown():
            startup_log.write("window_shown")
            window_ready.set()
            if startup_splash:
                startup_splash.stop()
            if ALWAYS_START_MAXIMIZED or (state and state.get("maximized")):
                window.maximize()
            if activation_pending.is_set():
                activate_process_window(os.getpid(), APP_NAME, logger=startup_log)

        def on_window_loaded():
            startup_log.write("window_loaded")

        closing_allowed = threading.Event()
        closing_flush_started = threading.Event()

        def show_close_error(text):
            ctypes.windll.user32.MessageBoxW(None, text, APP_NAME, 0x10)

        def flush_then_close():
            """Flush outside the native closing callback to avoid a WebView deadlock."""
            try:
                pending_data = window.evaluate_js(
                    "window.knizhnitsaBeforeClose ? window.knizhnitsaBeforeClose() : null"
                )
                if isinstance(pending_data, str) and pending_data.strip():
                    try:
                        closing_lang = (json.loads(pending_data).get("settings") or {}).get("lang", "ru")
                    except Exception:
                        closing_lang = settings.get("lang", "ru")
                    flush_result = api.flush_save(pending_data, closing_lang)
                    startup_log.write(
                        "close_flush",
                        success=bool(flush_result.get("ok")),
                        changed=bool(flush_result.get("changed")),
                        error=flush_result.get("error", ""),
                    )
                    if not flush_result.get("ok"):
                        show_close_error(
                            "Книжница не смогла сохранить последние изменения.\n\n"
                            "Окно останется открытым. Освободите место на диске или закройте программу, "
                            "которая мешает записи, и попробуйте ещё раз.\n\n"
                            + flush_result.get("error", "")
                        )
                        return
                else:
                    startup_log.write("close_flush", success=True, changed=False)
                save_window_state()
                closing_allowed.set()
                window.destroy()
            except Exception as exc:
                startup_log.write("close_flush", success=False, error=str(exc))
                show_close_error(
                    "Книжница не смогла проверить сохранение последних изменений.\n\n"
                    "Окно останется открытым. Попробуйте закрыть его ещё раз."
                )
            finally:
                if not closing_allowed.is_set():
                    closing_flush_started.clear()

        def on_window_closing(*_):
            startup_log.write("window_closing", allowed=closing_allowed.is_set())
            if closing_allowed.is_set():
                return
            if not closing_flush_started.is_set():
                closing_flush_started.set()
                threading.Thread(target=flush_then_close, name="close-flush", daemon=True).start()
            # Returning immediately is essential: synchronous browser calls
            # from the native closing callback can deadlock WebView2.
            return False

        window.events.shown += on_window_shown
        window.events.loaded += on_window_loaded
        # Запоминаем размер/позицию на случай, если позже отключим автозапуск развёрнутым.
        window.events.resized += lambda *a: save_window_state()
        window.events.moved += lambda *a: save_window_state()
        window.events.closing += on_window_closing

        startup_log.write("webview_start_begin")
        webview.start()
        startup_log.write("webview_start_done")
    except Exception as exc:
        startup_log.write("fatal_error", error=type(exc).__name__, message=str(exc))
        if startup_splash:
            startup_splash.stop()
        ctypes.windll.user32.MessageBoxW(
            None,
            f"Книжница не смогла запуститься.\n\nПодробности записаны сюда:\n{startup_log.path}",
            APP_NAME,
            0x10,
        )
        raise
    finally:
        if startup_splash:
            startup_splash.stop()
        if single_instance:
            single_instance.close()
        if startup_log:
            startup_log.write("shutdown")


if __name__ == "__main__":
    main()
